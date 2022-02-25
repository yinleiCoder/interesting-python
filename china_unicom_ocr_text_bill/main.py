import re
import os
from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from win32com.client import constants, gencache
from win32com import client
from paddleocr import PaddleOCR, draw_ocr

'''
【功能描述】OCR提取增值税发票图片的用户编号、电费年月、金额，并生转成PDF
这里的思路是先转成word、再转成pdf,可以考虑使用多线程/协程。
现阶段先写第一版本（quick start），暂不讲究编程规范

OCR: pip install paddlepaddle
     pip install paddleocr
[这个库有坑，官方没说，推荐版本用aconada安装python3.8]
注：(精准识别度高)https://github.com/PaddlePaddle/PaddleOCR/blob/release/2.4/README_ch.md
    需要手动下载中英文通用PP-OCR server模型（143.4M）官方的检测模型、方向分类器、识别模型的3个推理模型,将这3个推理模型压缩包放到电脑本地。
    也可以自己训练模型。
    
docx: pip install python-docx
pdf: pip install pywin32
'''

det_model_dir = "./inference/ch_ppocr_server_v2.0_det_infer/"
rec_model_dir = "./inference/ch_ppocr_server_v2.0_rec_infer/"
cls_model_dir = "./inference/ch_ppocr_mobile_v2.0_cls_infer/"


class ChinaUnicomBillPaddleOCR(PaddleOCR):
    def __init__(self, **kwargs):
        super(ChinaUnicomBillPaddleOCR, self).__init__(**kwargs)
        current_path = os.path.abspath(os.path.dirname(__file__))
        imgs_filename = os.listdir(f'{current_path}\\china_unicom_bill_imgs')
        self.all_imgs_path = [os.path.join(f'{current_path}\\china_unicom_bill_imgs\\', f) for f in imgs_filename]
        self.ocr_data = []

    def apply(self, det=True, rec=True, cls=True):
        for img in self.all_imgs_path:
            results = self.ocr(img, det=det, rec=rec, cls=cls, )
            money = []
            user_number=''
            elec_date=''
            for result in results:
                if result[1][0].startswith('用户编号'):
                    data = result[1][0]
                    temp_index = result[1][0].find('电')
                    user_number = data[0:temp_index-1]
                    elec_date = '2021年11月' + data[temp_index:temp_index+2]
                amounts = re.findall('([0-9]*?\.[0-9]{2})', result[1][0])
                # print(result[1][0])
                if amounts:
                    money.append(float(amounts[0]))
            self.ocr_data.append(f"{elec_date}_价税合计{money[-1]}元_{user_number}")
        print(self.ocr_data)

    def imgToDocx(self):
        for index, img_path in enumerate(self.all_imgs_path):
            document = Document()
            # 发票不适合纵向A4纸张，所以旋转一下
            section = document.sections[0]
            new_width, new_height = section.page_height, section.page_width
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height
            document.add_picture(f"{img_path}", width=Cm(21))
            document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            document.save(f'ocr_docx_result\\{self.ocr_data[index]}.docx')

    def docxToPDF(self):
        for index, img_path in enumerate(self.all_imgs_path):
            # 创建word程序对象
            # word = gencache.EnsureDispatch("Word.Application") # 需要先运行这句，才能初始化com组件，constants才不能为空
            word = client.DispatchEx("Word.Application")
            # 读取word文件
            doc = word.Documents.Open(f"E:\\PycharmProjects\\funnyPython\\china_unicom_ocr_text_bill\\ocr_docx_result\\{self.ocr_data[index]}.docx", ReadOnly=1)
            # word转换pdf(更多信息访问office开发人员中心文档)
            doc.ExportAsFixedFormat(f"E:\\PycharmProjects\\funnyPython\\china_unicom_ocr_text_bill\\ocr_pdf_result\\{self.ocr_data[index]}.pdf", constants.wdExportFormatPDF)
            word.Quit()

if __name__ == '__main__':
    # yl_ocr = ChinaUnicomBillPaddleOCR(cls_model_dir=cls_model_dir, det_model_dir=det_model_dir, rec_model_dir=rec_model_dir)
    yl_ocr = ChinaUnicomBillPaddleOCR()
    yl_ocr.apply()
    yl_ocr.imgToDocx()
    yl_ocr.docxToPDF()
    # 显示结果
    # from PIL import Image
    # image = Image.open(imgs).convert('RGB')
    # boxes = [line[0] for line in results]
    # txts = [line[1][0] for line in results]
    # scores = [line[1][1] for line in results]
    # im_show = draw_ocr(image, boxes, txts, scores, font_path='./fonts/simfang.ttf')
    # im_show = Image.fromarray(im_show)
    # im_show.save('result.jpg')
