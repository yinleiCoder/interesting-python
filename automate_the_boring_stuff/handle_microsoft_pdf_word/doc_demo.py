import docx
import win32com.client

"""
python-docx: pip install python-docx
pywin32: pip install pywin32
word转pdf需要pywin32，且电脑上装有Microsoft Word

"""
# doc = docx.Document('迎新春联通公司活动.docx')
# print(doc.paragraphs[10].text)
# print(doc.paragraphs[10].style)
# doc.paragraphs[10].style = 'Normal'
# doc.paragraphs[10].runs[0].style =
# doc.paragraphs[10].runs[0].underline = True
# doc.save('result.docx')

# 转word为pdf
wordFilename = 'helloworld.docx'
pdfFilename = 'hellopdf.pdf'
doc = docx.Document()
doc.save(wordFilename)
wdFormatPDF = 17
wordObj = win32com.client.Dispatch('Word.Application')
docObj = wordObj.Documents.Open(wordFilename)
docObj.SaveAs(pdfFilename, FileFormat=wdFormatPDF)
docObj.Close()
wordObj.Quit()

# doc = docx.Document()
# doc.add_paragraph('Hello, World!', 'Title')
# paraObj = doc.add_paragraph('i am yinlei')
# paraObj.add_run('!')
# doc.add_heading('Header 0', 0)
# doc.add_heading('Header 1', 1)
# doc.add_heading('Header 2', 2)
# doc.add_heading('Header 3', 3)
# doc.paragraphs[-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)# 换页符
# doc.add_paragraph('second page!')
# doc.add_picture('yinlei.jpg', width=docx.shared.Inches(1), height=docx.shared.Cm(4))
# doc.save("helloworld.docx")

# def getText(filename):
#     doc = docx.Document(filename)
#     fullText = []
#     for para in doc.paragraphs:
#         fullText.append(para.text)
#     return '\n'.join(fullText)
# print(getText('迎新春联通公司活动.docx'))# 取得完整的文本
# doc = docx.Document('迎新春联通公司活动.docx')
# print(len(doc.paragraphs))
# print(doc.paragraphs[0].text)
# print(doc.paragraphs[10].text)
# print(len(doc.paragraphs[10].runs))
# print(doc.paragraphs[10].runs[0].text)